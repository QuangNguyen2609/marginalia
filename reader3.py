"""
Parses book files (EPUB, PDF, DOCX, MOBI) into a structured object
that can be used to serve the book via a web interface.
"""

import os
import pickle
import shutil
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Any
from datetime import datetime
from urllib.parse import unquote

import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup, Comment

# --- Data structures ---

@dataclass
class ChapterContent:
    """
    Represents a physical file in the EPUB (Spine Item).
    A single file might contain multiple logical chapters (TOC entries).
    """
    id: str           # Internal ID (e.g., 'item_1')
    href: str         # Filename (e.g., 'part01.html')
    title: str        # Best guess title from file
    content: str      # Cleaned HTML with rewritten image paths
    text: str         # Plain text for search/LLM context
    order: int        # Linear reading order


@dataclass
class TOCEntry:
    """Represents a logical entry in the navigation sidebar."""
    title: str
    href: str         # original href (e.g., 'part01.html#chapter1')
    file_href: str    # just the filename (e.g., 'part01.html')
    anchor: str       # just the anchor (e.g., 'chapter1'), empty if none
    children: List['TOCEntry'] = field(default_factory=list)


@dataclass
class BookMetadata:
    """Metadata"""
    title: str
    language: str
    authors: List[str] = field(default_factory=list)
    description: Optional[str] = None
    publisher: Optional[str] = None
    date: Optional[str] = None
    identifiers: List[str] = field(default_factory=list)
    subjects: List[str] = field(default_factory=list)


@dataclass
class Book:
    """The Master Object to be pickled."""
    metadata: BookMetadata
    spine: List[ChapterContent]  # The actual content (linear files)
    toc: List[TOCEntry]          # The navigation tree
    images: Dict[str, str]       # Map: original_path -> local_path

    # Meta info
    source_file: str
    processed_at: str
    version: str = "3.0"


# --- Utilities ---

def clean_html_content(soup: BeautifulSoup) -> BeautifulSoup:

    # Remove dangerous/useless tags
    for tag in soup(['script', 'style', 'iframe', 'video', 'nav', 'form', 'button']):
        tag.decompose()

    # Remove HTML comments
    for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
        comment.extract()

    # Remove input tags
    for tag in soup.find_all('input'):
        tag.decompose()

    return soup


def extract_plain_text(soup: BeautifulSoup) -> str:
    """Extract clean text for LLM/Search usage."""
    text = soup.get_text(separator=' ')
    # Collapse whitespace
    return ' '.join(text.split())


def parse_toc_recursive(toc_list, depth=0) -> List[TOCEntry]:
    """
    Recursively parses the TOC structure from ebooklib.
    """
    result = []

    for item in toc_list:
        # ebooklib TOC items are either `Link` objects or tuples (Section, [Children])
        if isinstance(item, tuple):
            section, children = item
            entry = TOCEntry(
                title=section.title,
                href=section.href,
                file_href=section.href.split('#')[0],
                anchor=section.href.split('#')[1] if '#' in section.href else "",
                children=parse_toc_recursive(children, depth + 1)
            )
            result.append(entry)
        elif isinstance(item, epub.Link):
            entry = TOCEntry(
                title=item.title,
                href=item.href,
                file_href=item.href.split('#')[0],
                anchor=item.href.split('#')[1] if '#' in item.href else ""
            )
            result.append(entry)
        # Note: ebooklib sometimes returns direct Section objects without children
        elif isinstance(item, epub.Section):
             entry = TOCEntry(
                title=item.title,
                href=item.href,
                file_href=item.href.split('#')[0],
                anchor=item.href.split('#')[1] if '#' in item.href else ""
            )
             result.append(entry)

    return result


def get_fallback_toc(book_obj) -> List[TOCEntry]:
    """
    If TOC is missing, build a flat one from the Spine.
    """
    toc = []
    for item in book_obj.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            name = item.get_name()
            # Try to guess a title from the content or ID
            title = item.get_name().replace('.html', '').replace('.xhtml', '').replace('_', ' ').title()
            toc.append(TOCEntry(title=title, href=name, file_href=name, anchor=""))
    return toc


def extract_metadata_robust(book_obj) -> BookMetadata:
    """
    Extracts metadata handling both single and list values.
    """
    def get_list(key):
        data = book_obj.get_metadata('DC', key)
        return [x[0] for x in data] if data else []

    def get_one(key):
        data = book_obj.get_metadata('DC', key)
        return data[0][0] if data else None

    return BookMetadata(
        title=get_one('title') or "Untitled",
        language=get_one('language') or "en",
        authors=get_list('creator'),
        description=get_one('description'),
        publisher=get_one('publisher'),
        date=get_one('date'),
        identifiers=get_list('identifier'),
        subjects=get_list('subject')
    )


# --- Main Conversion Logic ---

def process_epub(epub_path: str, output_dir: str) -> Book:

    # 1. Load Book
    print(f"Loading {epub_path}...")
    book = epub.read_epub(epub_path)

    # 2. Extract Metadata
    metadata = extract_metadata_robust(book)

    # 3. Prepare Output Directories
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    images_dir = os.path.join(output_dir, 'images')
    os.makedirs(images_dir, exist_ok=True)

    # 4. Extract Images & Build Map
    print("Extracting images...")
    image_map = {} # Key: internal_path, Value: local_relative_path

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_IMAGE:
            # Normalize filename
            original_fname = os.path.basename(item.get_name())
            # Sanitize filename for OS
            safe_fname = "".join([c for c in original_fname if c.isalpha() or c.isdigit() or c in '._-']).strip()

            # Save to disk
            local_path = os.path.join(images_dir, safe_fname)
            with open(local_path, 'wb') as f:
                f.write(item.get_content())

            # Map keys: We try both the full internal path and just the basename
            # to be robust against messy HTML src attributes
            rel_path = f"images/{safe_fname}"
            image_map[item.get_name()] = rel_path
            image_map[original_fname] = rel_path

    # 5. Process TOC
    print("Parsing Table of Contents...")
    toc_structure = parse_toc_recursive(book.toc)
    if not toc_structure:
        print("Warning: Empty TOC, building fallback from Spine...")
        toc_structure = get_fallback_toc(book)

    # 6. Process Content (Spine-based to preserve HTML validity)
    print("Processing chapters...")
    spine_chapters = []

    # We iterate over the spine (linear reading order)
    for i, spine_item in enumerate(book.spine):
        item_id, linear = spine_item
        item = book.get_item_with_id(item_id)

        if not item:
            continue

        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            # Raw content
            raw_content = item.get_content().decode('utf-8', errors='ignore')
            soup = BeautifulSoup(raw_content, 'html.parser')

            # A. Fix Images
            for img in soup.find_all('img'):
                src = img.get('src', '')
                if not src: continue

                # Decode URL (part01/image%201.jpg -> part01/image 1.jpg)
                src_decoded = unquote(src)
                filename = os.path.basename(src_decoded)

                # Try to find in map
                if src_decoded in image_map:
                    img['src'] = image_map[src_decoded]
                elif filename in image_map:
                    img['src'] = image_map[filename]

            # B. Clean HTML
            soup = clean_html_content(soup)

            # C. Extract Body Content only
            body = soup.find('body')
            if body:
                # Extract inner HTML of body
                final_html = "".join([str(x) for x in body.contents])
            else:
                final_html = str(soup)

            # D. Create Object
            chapter = ChapterContent(
                id=item_id,
                href=item.get_name(), # Important: This links TOC to Content
                title=f"Section {i+1}", # Fallback, real titles come from TOC
                content=final_html,
                text=extract_plain_text(soup),
                order=i
            )
            spine_chapters.append(chapter)

    # 7. Final Assembly
    final_book = Book(
        metadata=metadata,
        spine=spine_chapters,
        toc=toc_structure,
        images=image_map,
        source_file=os.path.basename(epub_path),
        processed_at=datetime.now().isoformat()
    )

    return final_book


# --- PDF Processing ---

def process_pdf(pdf_path: str, output_dir: str) -> Book:
    """Process a PDF file using PyMuPDF (fitz)."""
    import fitz

    print(f"Loading PDF: {pdf_path}...")
    doc = fitz.open(pdf_path)

    # Prepare output
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    images_dir = os.path.join(output_dir, 'images')
    os.makedirs(images_dir, exist_ok=True)

    # Metadata
    meta = doc.metadata or {}
    metadata = BookMetadata(
        title=meta.get("title") or os.path.splitext(os.path.basename(pdf_path))[0],
        language="en",
        authors=[meta.get("author")] if meta.get("author") else [],
        description=meta.get("subject"),
        publisher=meta.get("producer"),
        date=meta.get("creationDate"),
    )

    # TOC
    raw_toc = doc.get_toc()
    toc_entries = _build_pdf_toc(raw_toc) if raw_toc else []

    # Copy source PDF to output dir so it can be served for PDF.js rendering
    import shutil as _shutil
    source_pdf_dest = os.path.join(output_dir, 'source.pdf')
    _shutil.copy2(pdf_path, source_pdf_dest)

    # Create one chapter entry per page; content is a sized placeholder.
    # PDF.js in the browser will render each page into a canvas with a text layer.
    image_map = {}
    chapters = []
    for i, page in enumerate(doc):
        plain_text = page.get_text()
        w, h = page.rect.width, page.rect.height

        # aspect-ratio reserves the correct space before JS renders the canvas,
        # keeping TOC scroll positions accurate.
        content = (
            f'<div class="pdf-page-placeholder" data-page="{i+1}" '
            f'style="aspect-ratio:{w}/{h}"></div>'
        )

        chapter = ChapterContent(
            id=f"page_{i+1}",
            href=f"page_{i+1}.html",
            title=f"Page {i+1}",
            content=content,
            text=' '.join(plain_text.split()),
            order=i
        )
        chapters.append(chapter)

    # If no TOC, build flat one from pages
    if not toc_entries:
        toc_entries = [
            TOCEntry(title=f"Page {i+1}", href=f"page_{i+1}.html", file_href=f"page_{i+1}.html", anchor="")
            for i in range(len(chapters))
        ]

    doc.close()

    return Book(
        metadata=metadata,
        spine=chapters,
        toc=toc_entries,
        images=image_map,
        source_file=os.path.basename(pdf_path),
        processed_at=datetime.now().isoformat()
    )


def _build_pdf_toc(raw_toc: list) -> List[TOCEntry]:
    """Convert PyMuPDF TOC [(level, title, page_num), ...] into TOCEntry tree."""
    root = []
    stack = [(0, root)]  # (level, children_list)

    for level, title, page_num in raw_toc:
        entry = TOCEntry(
            title=title,
            href=f"page_{page_num}.html",
            file_href=f"page_{page_num}.html",
            anchor=""
        )
        # Find correct parent
        while len(stack) > 1 and stack[-1][0] >= level:
            stack.pop()

        stack[-1][1].append(entry)
        stack.append((level, entry.children))

    return root


# --- DOCX Processing ---

def process_docx(docx_path: str, output_dir: str) -> Book:
    """Process a DOCX file using python-docx."""
    ext = os.path.splitext(docx_path)[1].lower()
    if ext == '.doc':
        raise ValueError(
            "Legacy .doc format is not supported. "
            "Please convert to .docx using LibreOffice or Microsoft Word."
        )

    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    print(f"Loading DOCX: {docx_path}...")
    doc = Document(docx_path)

    # Prepare output
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    images_dir = os.path.join(output_dir, 'images')
    os.makedirs(images_dir, exist_ok=True)

    # Metadata
    props = doc.core_properties
    metadata = BookMetadata(
        title=props.title or os.path.splitext(os.path.basename(docx_path))[0],
        language="en",
        authors=[props.author] if props.author else [],
        description=props.subject,
    )

    # Extract images
    image_map = {}
    img_counter = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img_ext = os.path.splitext(rel.target_ref)[1] or ".png"
                img_name = f"img_{img_counter}{img_ext}"
                img_path = os.path.join(images_dir, img_name)
                with open(img_path, "wb") as f:
                    f.write(img_data)
                image_map[rel.target_ref] = f"images/{img_name}"
                img_counter += 1
            except Exception:
                continue

    # Split into chapters by headings
    chapters = []
    toc_entries = []
    current_html_parts = []
    current_title = "Introduction"
    chapter_idx = 0

    heading_styles = {'Heading 1', 'Heading 2', 'Heading 3',
                      'Heading1', 'Heading2', 'Heading3'}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        if style_name in heading_styles and para.text.strip():
            # Save previous chapter if it has content
            if current_html_parts:
                html_str = "\n".join(current_html_parts)
                soup = BeautifulSoup(html_str, 'html.parser')
                chapters.append(ChapterContent(
                    id=f"chapter_{chapter_idx}",
                    href=f"chapter_{chapter_idx}.html",
                    title=current_title,
                    content=html_str,
                    text=extract_plain_text(soup),
                    order=chapter_idx
                ))
                toc_entries.append(TOCEntry(
                    title=current_title,
                    href=f"chapter_{chapter_idx}.html",
                    file_href=f"chapter_{chapter_idx}.html",
                    anchor=""
                ))
                chapter_idx += 1
                current_html_parts = []

            current_title = para.text.strip()

        # Convert paragraph to HTML
        tag = "p"
        if "Heading 1" in style_name or "Heading1" in style_name:
            tag = "h1"
        elif "Heading 2" in style_name or "Heading2" in style_name:
            tag = "h2"
        elif "Heading 3" in style_name or "Heading3" in style_name:
            tag = "h3"

        text = para.text.strip()
        if text:
            current_html_parts.append(f"<{tag}>{_escape_html(text)}</{tag}>")

    # Save final chapter
    if current_html_parts:
        html_str = "\n".join(current_html_parts)
        soup = BeautifulSoup(html_str, 'html.parser')
        chapters.append(ChapterContent(
            id=f"chapter_{chapter_idx}",
            href=f"chapter_{chapter_idx}.html",
            title=current_title,
            content=html_str,
            text=extract_plain_text(soup),
            order=chapter_idx
        ))
        toc_entries.append(TOCEntry(
            title=current_title,
            href=f"chapter_{chapter_idx}.html",
            file_href=f"chapter_{chapter_idx}.html",
            anchor=""
        ))

    if not chapters:
        # Fallback: single chapter with all content
        all_text = "\n".join(f"<p>{_escape_html(p.text)}</p>" for p in doc.paragraphs if p.text.strip())
        soup = BeautifulSoup(all_text, 'html.parser')
        chapters.append(ChapterContent(
            id="chapter_0", href="chapter_0.html", title="Content",
            content=all_text, text=extract_plain_text(soup), order=0
        ))
        toc_entries.append(TOCEntry(title="Content", href="chapter_0.html", file_href="chapter_0.html", anchor=""))

    return Book(
        metadata=metadata,
        spine=chapters,
        toc=toc_entries,
        images=image_map,
        source_file=os.path.basename(docx_path),
        processed_at=datetime.now().isoformat()
    )


def _escape_html(text: str) -> str:
    """Basic HTML escaping."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# --- MOBI Processing ---

def process_mobi(mobi_path: str, output_dir: str) -> Book:
    """Process a MOBI file."""
    import mobi

    print(f"Loading MOBI: {mobi_path}...")

    # Prepare output
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    images_dir = os.path.join(output_dir, 'images')
    os.makedirs(images_dir, exist_ok=True)

    # Extract MOBI content
    tempdir, filepath = mobi.extract(mobi_path)

    try:
        # Read extracted HTML
        with open(filepath, 'r', errors='ignore') as f:
            raw_html = f.read()

        soup = BeautifulSoup(raw_html, 'html.parser')
        soup = clean_html_content(soup)

        # Try to split by headings
        body = soup.find('body') or soup
        chapters = []
        toc_entries = []

        headings = body.find_all(['h1', 'h2', 'h3'])

        if headings:
            # Split content at headings
            all_elements = list(body.children)
            current_parts = []
            current_title = "Introduction"
            idx = 0

            for elem in all_elements:
                if hasattr(elem, 'name') and elem.name in ('h1', 'h2', 'h3') and elem.get_text(strip=True):
                    # Save previous chapter
                    if current_parts:
                        html_str = "".join(str(e) for e in current_parts)
                        text_soup = BeautifulSoup(html_str, 'html.parser')
                        chapters.append(ChapterContent(
                            id=f"chapter_{idx}", href=f"chapter_{idx}.html",
                            title=current_title, content=html_str,
                            text=extract_plain_text(text_soup), order=idx
                        ))
                        toc_entries.append(TOCEntry(
                            title=current_title, href=f"chapter_{idx}.html",
                            file_href=f"chapter_{idx}.html", anchor=""
                        ))
                        idx += 1
                        current_parts = []
                    current_title = elem.get_text(strip=True)

                current_parts.append(elem)

            # Save final chapter
            if current_parts:
                html_str = "".join(str(e) for e in current_parts)
                text_soup = BeautifulSoup(html_str, 'html.parser')
                chapters.append(ChapterContent(
                    id=f"chapter_{idx}", href=f"chapter_{idx}.html",
                    title=current_title, content=html_str,
                    text=extract_plain_text(text_soup), order=idx
                ))
                toc_entries.append(TOCEntry(
                    title=current_title, href=f"chapter_{idx}.html",
                    file_href=f"chapter_{idx}.html", anchor=""
                ))
        else:
            # Single chapter fallback
            html_str = str(body)
            chapters.append(ChapterContent(
                id="chapter_0", href="chapter_0.html",
                title="Content", content=html_str,
                text=extract_plain_text(soup), order=0
            ))
            toc_entries.append(TOCEntry(
                title="Content", href="chapter_0.html",
                file_href="chapter_0.html", anchor=""
            ))

        # Extract title from HTML
        title_tag = soup.find('title')
        title = title_tag.get_text(strip=True) if title_tag else os.path.splitext(os.path.basename(mobi_path))[0]

        metadata = BookMetadata(title=title, language="en")

    finally:
        # Clean up temp directory
        shutil.rmtree(tempdir, ignore_errors=True)

    return Book(
        metadata=metadata,
        spine=chapters,
        toc=toc_entries,
        images={},
        source_file=os.path.basename(mobi_path),
        processed_at=datetime.now().isoformat()
    )


# --- Format Router ---

def process_book(file_path: str, output_dir: str) -> Book:
    """Route to the appropriate processor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    processors = {
        '.epub': process_epub,
        '.pdf': process_pdf,
        '.mobi': process_mobi,
        '.docx': process_docx,
        '.doc': process_docx,
    }

    processor = processors.get(ext)
    if not processor:
        raise ValueError(f"Unsupported format: {ext}")

    return processor(file_path, output_dir)


def save_to_pickle(book: Book, output_dir: str):
    p_path = os.path.join(output_dir, 'book.pkl')
    with open(p_path, 'wb') as f:
        pickle.dump(book, f)
    print(f"Saved structured data to {p_path}")


# --- CLI ---

if __name__ == "__main__":

    import sys
    if len(sys.argv) < 2:
        print("Usage: python reader3.py <file>")
        print("Supported formats: .epub, .pdf, .docx, .mobi")
        sys.exit(1)

    input_file = sys.argv[1]
    assert os.path.exists(input_file), "File not found."
    out_dir = os.path.splitext(input_file)[0] + "_data"

    book_obj = process_book(input_file, out_dir)
    save_to_pickle(book_obj, out_dir)
    print("\n--- Summary ---")
    print(f"Title: {book_obj.metadata.title}")
    print(f"Authors: {', '.join(book_obj.metadata.authors)}")
    print(f"Physical Files (Spine): {len(book_obj.spine)}")
    print(f"TOC Root Items: {len(book_obj.toc)}")
    print(f"Images extracted: {len(book_obj.images)}")
